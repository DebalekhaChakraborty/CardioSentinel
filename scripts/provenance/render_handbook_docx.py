"""Render the v1.2 handbook markdown into .docx using v1.1 as the style template."""
import re, sys
from docx import Document
from docx.shared import Pt

SRC, TPL, OUT = sys.argv[1], sys.argv[2], sys.argv[3]
doc = Document(TPL)
body = doc.element.body
for el in list(body):                      # clear template content, keep styles/theme/sectPr
    if el.tag.endswith('}sectPr'):
        continue
    body.remove(el)

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)')

def add_runs(p, text):
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)       # links -> label
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)

def para(text, style=None):
    p = doc.add_paragraph(style=style)
    add_runs(p, text)
    return p

lines = open(SRC, encoding='utf-8').read().splitlines()
i, first_h1 = 0, True
tables = codeblocks = 0
while i < len(lines):
    ln = lines[i]

    if ln.startswith('```'):                                   # code block
        i += 1; buf = []
        while i < len(lines) and not lines[i].startswith('```'):
            buf.append(lines[i]); i += 1
        i += 1; codeblocks += 1
        for b in buf:
            p = doc.add_paragraph(style='CodeBlock')
            p.add_run(b if b.strip() else ' ')
        continue

    if ln.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s:|-]+\|$', lines[i+1]):
        rows = []
        while i < len(lines) and lines[i].startswith('|'):
            if not re.match(r'^\|[\s:|-]+\|$', lines[i]):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
            i += 1
        ncol = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=ncol); t.style = 'Table Grid'; tables += 1
        for ri, r in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(ncol):
                txt = r[ci] if ci < len(r) else ''
                cp = cells[ci].paragraphs[0]
                add_runs(cp, txt)
                if ri == 0:
                    for run in cp.runs:
                        run.bold = True
        continue

    if ln.startswith('#'):
        lvl = len(ln) - len(ln.lstrip('#'))
        txt = ln.lstrip('#').strip()
        if lvl == 1 and first_h1:
            doc.add_paragraph(txt, style='Title'); first_h1 = False
        else:
            para(txt, style='Heading1' if lvl <= 2 else 'Heading2')
        i += 1; continue

    if ln.strip() in ('---', '***'):
        i += 1; continue

    m = re.match(r'^(\s*)[-*]\s+(.*)', ln)
    if m:
        para(m.group(2), style='ListBullet'); i += 1; continue
    m = re.match(r'^(\s*)\d+\.\s+(.*)', ln)
    if m:
        para(m.group(2), style='ListNumber'); i += 1; continue

    if ln.startswith('>'):
        p = para(ln.lstrip('> ').strip())
        p.paragraph_format.left_indent = Pt(24)
        for r in p.runs: r.italic = True
        i += 1; continue

    if ln.strip():                                             # wrap soft-wrapped prose
        buf = [ln.strip()]; i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#|\||>|```|\s*[-*]\s|\s*\d+\.\s|---$)', lines[i]):
            buf.append(lines[i].strip()); i += 1
        para(' '.join(buf))
        continue
    i += 1

doc.save(OUT)
print(f"  wrote {OUT}: {len(doc.element.body)} body elements, {tables} tables, {codeblocks} code blocks")
